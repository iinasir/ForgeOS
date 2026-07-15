from pathlib import Path

from docx import Document

from app.models.resume import Resume


class DOCXRenderer:
    """
    Export Resume to Microsoft Word (.docx)
    """

    @staticmethod
    def export(
        resume: Resume,
        output_path: str
    ) -> None:

        document = Document()

        # ------------------------
        # Personal Information
        # ------------------------

        document.add_heading(
            resume.personal.full_name,
            level=0
        )

        contact = (
            f"{resume.personal.email} | "
            f"{resume.personal.phone}"
        )

        if resume.personal.location:
            contact += f" | {resume.personal.location}"

        document.add_paragraph(contact)

        # ------------------------
        # Summary
        # ------------------------

        if resume.personal.summary:

            document.add_heading(
                "Professional Summary",
                level=1
            )

            document.add_paragraph(
                resume.personal.summary
            )

        # ------------------------
        # Skills
        # ------------------------

        if resume.skills:

            document.add_heading(
                "Skills",
                level=1
            )

            for skill in resume.skills:

                level = (
                    f" ({skill.level})"
                    if getattr(skill, "level", None)
                    else ""
                )

                document.add_paragraph(
                    f"{skill.name}{level}",
                    style="List Bullet"
                )

        # ------------------------
        # Experience
        # ------------------------

        if getattr(resume, "experiences", None):

            document.add_heading(
                "Experience",
                level=1
            )

            for exp in resume.experiences:

                title = (
                    f"{exp.position} - {exp.company}"
                )

                document.add_heading(
                    title,
                    level=2
                )

                document.add_paragraph(
                    f"{exp.start_date} - {exp.end_date}"
                )

                if exp.description:

                     for item in exp.description:

                         document.add_paragraph(
                             item,
                             style="List Bullet"
                         )

        # ------------------------
        # Education
        # ------------------------

        if getattr(resume, "education", None):

            document.add_heading(
                "Education",
                level=1
            )

            for edu in resume.education:

                text = (
                    f"{edu.degree} - "
                    f"{edu.institution}"
                )

                document.add_paragraph(
                    text,
                    style="List Bullet"
                )

        # ------------------------
        # Projects
        # ------------------------

        if getattr(resume, "projects", None):

            document.add_heading(
                "Projects",
                level=1
            )

            for project in resume.projects:

                document.add_heading(
                    project.name,
                    level=2
                )

                if getattr(project, "description", None):
                    document.add_paragraph(
                        project.description
                    )

        # ------------------------
        # Save
        # ------------------------

        output = Path(output_path)

        output.parent.mkdir(
            parents=True,
            exist_ok=True
        )

        document.save(str(output))